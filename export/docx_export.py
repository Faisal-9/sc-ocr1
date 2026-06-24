import io
from typing import Optional

from docx import Document


def export_to_docx(text: str, title: Optional[str] = "OCR Extracted Text") -> bytes:
    doc = Document()

    if title:
        doc.add_heading(title, level=1)

    for para in text.split("\n"):
        doc.add_paragraph(para)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()